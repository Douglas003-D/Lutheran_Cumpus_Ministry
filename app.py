import os
import math
import logging
import base64
import requests
import re
import csv
import io  # Standard io import for BytesIO and StringIO
from io import BytesIO, StringIO
from datetime import date, datetime, timedelta
from functools import wraps

from flask import Flask, render_template, request, redirect, url_for, session, flash, \
                  send_from_directory, send_file, jsonify, make_response
from flask_mysqldb import MySQL
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from flask_mailman import Mail, EmailMessage 
from itsdangerous import URLSafeTimedSerializer, SignatureExpired, BadTimeSignature
import MySQLdb.cursors

# Word Document Imports
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH # Added this to fix your alignment error

from config import Config
# --- Logging Configuration ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config.from_object(Config)

# --- Security & Session Settings ---
app.permanent_session_lifetime = timedelta(minutes=60)
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

# FILE SIZE LIMIT: 100MB
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024 

# --- Database & Mail Setup ---
mysql = MySQL(app)
mail = Mail(app)
def is_strong_password(password):
    """Checks if password meets: 8+ chars, 1 Upper, 1 Lower, 1 Digit, 1 Symbol."""
    if len(password) < 8:
        return False, "Password must be at least 8 characters long."
    if not re.search(r"[A-Z]", password):
        return False, "Password must contain at least one uppercase letter."
    if not re.search(r"[a-z]", password):
        return False, "Password must contain at least one lowercase letter."
    if not re.search(r"\d", password):
        return False, "Password must contain at least one number."
    if not re.search(r"[\W_]", password):
        return False, "Password must contain at least one special character (e.g. @, #, $, etc.)."
    return True, ""

# Helper to get serializer with the current secret key
def get_serializer():
    return URLSafeTimedSerializer(app.config['SECRET_KEY'])

# --- File Upload Configuration ---
UPLOAD_FOLDER = os.path.join(app.root_path, 'static', 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'jpg', 'png', 'jpeg', 'zip', 'mp4', 'webm', 'mov'}

# --- M-PESA HELPERS ---

def get_mpesa_access_token():
    """Fetches the OAuth2 token from Safaricom."""
    url = f"{app.config['MPESA_BASE_URL']}/oauth/v1/generate?grant_type=client_credentials"
    try:
        response = requests.get(url, auth=(app.config['MPESA_CONSUMER_KEY'], app.config['MPESA_CONSUMER_SECRET']), timeout=10)
        return response.json().get('access_token')
    except Exception as e:
        logger.error(f"Mpesa Auth Error: {e}")
        return None

def generate_stk_password():
    """Generates the dynamic password required for the STK Push."""
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    data_to_encode = app.config['MPESA_SHORTCODE'] + app.config['MPESA_PASSKEY'] + timestamp
    password = base64.b64encode(data_to_encode.encode()).decode('utf-8')
    return password, timestamp

# --- Generic Helpers ---

def allowed_file(filename):
    """Checks if the file extension is allowed."""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def delete_physical_file(filename):
    """Removes files from the static/uploads directory."""
    if filename:
        try:
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            logger.error(f"Error deleting physical file {filename}: {e}")

def get_file_size(filename):
    """Returns a human-readable file size string."""
    try:
        path = os.path.join(UPLOAD_FOLDER, filename)
        if os.path.exists(path):
            size_bytes = os.path.getsize(path)
            if size_bytes == 0: return "0B"
            size_name = ("B", "KB", "MB", "GB", "TB")
            i = int(math.floor(math.log(size_bytes, 1024)))
            p = math.pow(1024, i)
            val = round(size_bytes / p, 2)
            return f"{val} {size_name[i]}"
    except Exception as e:
        logger.error(f"Error calculating file size for {filename}: {e}")
    return "Unknown"

def admin_required(f):
    """Decorator to protect admin routes."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('loggedin') or session.get('role') != 'admin':
            flash("Unauthorized Access. Admin login required.", "danger")
            return redirect(url_for('admin_login'))
        return f(*args, **kwargs)
    return decorated_function

@app.context_processor
def inject_global_vars():
    """Injects site settings and unread counts into all templates."""
    settings = {}
    unread_messages = 0
    try:
        cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cur.execute("SELECT * FROM site_settings WHERE id=1")
        settings = cur.fetchone()
        
        cur.execute("SELECT COUNT(*) as count FROM contact_messages WHERE is_read=0")
        msg_res = cur.fetchone()
        if msg_res:
            unread_messages = msg_res['count']

        if not settings:
            cur.execute("INSERT INTO site_settings (id, office_location) VALUES (1, 'Main Campus')")
            mysql.connection.commit()
            cur.execute("SELECT * FROM site_settings WHERE id=1")
            settings = cur.fetchone()
        cur.close()
    except Exception as e:
        logger.error(f"Global context error: {e}")
    
    return dict(settings=settings or {}, unread_count=unread_messages, now=datetime.now())

# --- Public Routes ---

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/contacts')
def contacts():
    return render_template('contacts.html')

@app.route('/submit_contact', methods=['POST'])
def submit_contact():
    full_name = request.form.get('full_name')
    email = request.form.get('email')
    phone = request.form.get('phone')
    inquiry = request.form.get('inquiry_type')
    message = request.form.get('message')

    if not all([full_name, email, message]):
        flash("Please fill in all required fields.", "warning")
        return redirect(url_for('contacts'))

    try:
        cur = mysql.connection.cursor()
        cur.execute("""
            INSERT INTO contact_messages (full_name, email, phone, inquiry_type, message) 
            VALUES (%s, %s, %s, %s, %s)
        """, (full_name, email, phone, inquiry, message))
        mysql.connection.commit()
        cur.close()
        flash("Thank you! Your message has been received.", "success")
    except Exception as e:
        logger.error(f"Contact form error: {e}")
        flash("Sorry, something went wrong. Please try again later.", "danger")
    
    return redirect(url_for('contacts'))
# --- GALLERY ROUTES ---


@app.route('/donate')
def donate():
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cur.execute("""
        SELECT paybill_no, paybill_acc_no, bank_name, account_no, account_name, paypal_email, mpesa_name 
        FROM site_settings WHERE id=1
    """)
    finance = cur.fetchone()
    cur.close()
    return render_template('donate.html', finance=finance)

@app.route('/initiate-donation', methods=['POST'])
def initiate_donation():
    try:
        data = request.get_json()
        raw_phone = data.get('phone') 
        amount = data.get('amount')

        if not raw_phone or not amount:
            return jsonify({'status': 'error', 'message': 'Missing phone or amount'}), 400

        # --- PHONE CLEANER ---
        # Ensures format is 2547XXXXXXXX or 2541XXXXXXXX
        phone = str(raw_phone).strip().replace("+", "")
        if phone.startswith("0"):
            phone = "254" + phone[1:]
        elif phone.startswith("7") or phone.startswith("1"):
            phone = "254" + phone
        # ---------------------

        token = get_mpesa_access_token()
        password, timestamp = generate_stk_password()
        
        if not token:
            print("!!! MPESA AUTH ERROR: Could not fetch access token !!!")
            return jsonify({'status': 'error', 'message': 'M-Pesa authentication failed.'}), 500

        headers = {"Authorization": f"Bearer {token}"}
        
        request_body = {
            "BusinessShortCode": app.config['MPESA_SHORTCODE'],
            "Password": password,
            "Timestamp": timestamp,
            "TransactionType": "CustomerPayBillOnline", 
            "Amount": int(float(amount)), 
            "PartyA": phone,
            "PartyB": app.config['MPESA_SHORTCODE'],
            "PhoneNumber": phone,
            "CallBackURL": app.config['MPESA_CALLBACK_URL'],
            "AccountReference": "LCM Donation",
            "TransactionDesc": "Ministry Support"
        }

        api_url = f"{app.config['MPESA_BASE_URL']}/mpesa/stkpush/v1/processrequest"
        response = requests.post(api_url, json=request_body, headers=headers, timeout=30)
        res_data = response.json()

        # FORCE DEBUG PRINT TO TERMINAL
        print("\n--- SAFARICOM DEBUG START ---")
        print(f"Status: {response.status_code}")
        print(f"Response: {res_data}")
        print("--- SAFARICOM DEBUG END ---\n")
        
        return jsonify(res_data)

    except Exception as e:
        print(f"CRITICAL SYSTEM ERROR: {str(e)}")
        return jsonify({'status': 'error', 'message': "Internal server error occurred."}), 500
  
@app.route('/mpesa-callback', methods=['POST'])
def mpesa_callback():
    """Handles the M-Pesa response and saves successful transactions to MySQL."""
    data = request.get_json()
    
    # Print the raw data to see the structure in the terminal
    print(f"\n[CALLBACK RECEIVED]: {data}\n")
    
    try:
        stk_callback = data['Body']['stkCallback']
        result_code = stk_callback['ResultCode']
        result_desc = stk_callback['ResultDesc']
        checkout_request_id = stk_callback['CheckoutRequestID']

        if result_code == 0:
            # Payment was successful
            items = stk_callback['CallbackMetadata']['Item']
            
            # Extract specific values from the metadata list
            mpesa_receipt = next((item['Value'] for item in items if item['Name'] == 'MpesaReceiptNumber'), "N/A")
            amount_paid = next((item['Value'] for item in items if item['Name'] == 'Amount'), 0)
            phone_number = next((item['Value'] for item in items if item['Name'] == 'PhoneNumber'), "N/A")

            print(f"SUCCESS: {mpesa_receipt} | Amount: {amount_paid} | Phone: {phone_number}")

            # --- SAVE TO MYSQL ---
            cur = mysql.connection.cursor()
            try:
                cur.execute("""
                    INSERT INTO donations (
                        receipt_number, 
                        amount, 
                        phone_number, 
                        checkout_request_id, 
                        status, 
                        created_at
                    ) VALUES (%s, %s, %s, %s, %s, %s)
                """, (
                    mpesa_receipt, 
                    amount_paid, 
                    phone_number, 
                    checkout_request_id, 
                    'Completed', 
                    datetime.now()
                ))
                mysql.connection.commit()
                print(f"DB Update: Transaction {mpesa_receipt} saved successfully.")
            except Exception as db_err:
                print(f"DATABASE ERROR: {db_err}")
            finally:
                cur.close()

        else:
            # ResultCode != 0 means user cancelled or transaction failed
            print(f"PAYMENT FAILED: {result_desc} (Code: {result_code})")

    except Exception as e:
        print(f"CALLBACK PROCESSING ERROR: {e}")

    # Safaricom requires this specific JSON response
    return jsonify({"ResultCode": 0, "ResultDesc": "Accepted"})

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        full_name = request.form.get('name')
        email = request.form.get('email')
        phone = request.form.get('phone')
        campus = request.form.get('campus') 
        
        if not all([full_name, email]):
            flash("Please provide at least a name and email.", "warning")
            return redirect(url_for('register'))

        cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        try:
            cur.execute("SELECT id FROM students WHERE email = %s", (email,))
            if cur.fetchone():
                flash("This email is already registered in our records.", "info")
                return redirect(url_for('index'))

            cur.execute("""
                INSERT INTO students (full_name, email, phone, school_campus) 
                VALUES (%s, %s, %s, %s)
            """, (full_name, email, phone, campus))
            mysql.connection.commit()
            flash(f"Welcome {full_name}! Registration successful.", "success")
            return redirect(url_for('index'))
        except Exception as e:
            mysql.connection.rollback()
            logger.error(f"Registration Error: {e}")
            flash(f"Database error: {str(e)}", "danger")
        finally:
            cur.close()
    return render_template('register.html')

@app.route('/events')
def events():
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    # Remove the WHERE clause to see everything in the database
    cur.execute("SELECT * FROM events ORDER BY event_date DESC") 
    data = cur.fetchall()
    cur.close()
    return render_template('events.html', events=data)
@app.route('/downloads')
def downloads():
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cur.execute("SELECT id, display_name, filename, uploaded_at FROM downloads ORDER BY uploaded_at DESC")
    files = cur.fetchall()
    cur.close()
    for f in files: 
        f['size'] = get_file_size(f['filename'])
    return render_template('downloads.html', files=files)

# --- Admin Auth Routes ---

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cur.execute("SELECT COUNT(*) as count FROM admin")
    if cur.fetchone()['count'] == 0:
        cur.close()
        return redirect(url_for('admin_setup'))

    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        cur.execute("SELECT * FROM admin WHERE username=%s", (username,))
        user = cur.fetchone()
        cur.close()
        
        if user and check_password_hash(user['password'], password):
            session.permanent = True
            session.update({
                'loggedin': True, 
                'role': 'admin', 
                'username': user['username'],
                'admin_id': user['id'],
                'email': user['email']
            })
            flash("Access Granted. Welcome back.", "success")
            return redirect(url_for('admin_dashboard'))
        flash("Invalid Credentials provided.", "danger")
    return render_template('admin_login.html')

@app.route('/admin/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form.get('email')
        cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cur.execute("SELECT * FROM admin WHERE email=%s", (email,))
        user = cur.fetchone()
        cur.close()

        if user:
            # Generate the timed token
            s = get_serializer()
            token = s.dumps(email, salt='password-reset-salt')
            
            # Generate the absolute URL for the email link
            link = url_for('reset_with_token', token=token, _external=True)
            
            try:
                msg = EmailMessage(
                    "Admin Password Reset Request",
                    f"To reset your password, visit the following link:\n{link}\n\n"
                    f"This link will expire in 30 minutes.",
                    to=[email]
                )
                msg.send()
                flash("A reset link has been sent to your email address.", "info")
            except Exception as e:
                app.logger.error(f"Mail delivery failed: {e}")
                flash("Internal mail server error. Please try again later.", "danger")
        else:
            # We use the same message for security to prevent email harvesting
            flash("If that email is in our system, a reset link has been sent.", "info")
        
        return redirect(url_for('admin_login'))
        
    return render_template('admin_forgot.html')


# --- RESET PASSWORD FINAL STEP ---
@app.route('/admin/reset_password/<token>', methods=['GET', 'POST'])
def reset_with_token(token):
    s = get_serializer()
    try:
        # Load the email from the token; expires in 1800 seconds (30 mins)
        email = s.loads(token, salt='password-reset-salt', max_age=1800)
    except SignatureExpired:
        flash("The reset link has expired. Please request a new one.", "danger")
        return redirect(url_for('forgot_password'))
    except BadTimeSignature:
        flash("The reset link is invalid.", "danger")
        return redirect(url_for('forgot_password'))
    except Exception:
        flash("An error occurred during verification.", "danger")
        return redirect(url_for('forgot_password'))

    if request.method == 'POST':
        new_pass = request.form.get('new_password')
        confirm = request.form.get('confirm_password')
        
        if new_pass != confirm:
            flash("Passwords do not match. Please try again.", "danger")
            return render_template('admin_reset_new.html', token=token)

        # Hash and Update MySQL
        hashed_pw = generate_password_hash(new_pass)
        cur = mysql.connection.cursor()
        cur.execute("UPDATE admin SET password=%s WHERE email=%s", (hashed_pw, email))
        mysql.connection.commit()
        cur.close()
        
        flash("Your password has been successfully updated!", "success")
        return redirect(url_for('admin_login'))
        
    # Render the form on GET request
    return render_template('admin_reset_new.html', token=token)

# STEP 1: Request the change
@app.route('/admin/change_request', methods=['GET', 'POST'])
@admin_required
def admin_change_request():
    if request.method == 'POST':
        email = request.form.get('email')
        
        if email != session.get('email'):
            flash("The email entered does not match your logged-in account.", "danger")
            return redirect(url_for('admin_change_request'))

        s = get_serializer()
        token = s.dumps(email, salt='secure-change-salt')
        link = url_for('admin_confirm_change', token=token, _external=True)

        try:
            msg = EmailMessage(
                "Security: Password Change Request",
                f"Use this link to change your password: {link}",
                to=[email]
            )
            msg.send()
            flash("Check your email for the verification link.", "info")
            return redirect(url_for('admin_dashboard'))
        except Exception:
            flash("Failed to send email.", "danger")

    return render_template('admin_change_request.html')

# STEP 2: The actual update (The link from the email)
@app.route('/admin/confirm_change/<token>', methods=['GET', 'POST'])
def admin_confirm_change(token):
    s = get_serializer()
    try:
        email = s.loads(token, salt='secure-change-salt', max_age=900)
    except:
        flash("The link has expired or is invalid.", "danger")
        return redirect(url_for('admin_change_request'))

    if request.method == 'POST':
        old_pw = request.form.get('old_password')
        new_pw = request.form.get('new_password')
        conf_pw = request.form.get('confirm_password')

        cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cur.execute("SELECT password FROM admin WHERE email=%s", (email,))
        user = cur.fetchone()

        if user and check_password_hash(user['password'], old_pw):
            if new_pw == conf_pw:
                hashed = generate_password_hash(new_pw)
                cur.execute("UPDATE admin SET password=%s WHERE email=%s", (hashed, email))
                mysql.connection.commit()
                cur.close()
                flash("Success! Password updated.", "success")
                return redirect(url_for('admin_dashboard'))
            else:
                flash("New passwords do not match.", "warning")
        else:
            flash("Incorrect current password.", "danger")
        cur.close()

    return render_template('admin_change_confirm.html', token=token)

@app.route('/admin/mark_read/<int:id>')
@admin_required
def mark_read(id):
    cur = mysql.connection.cursor()
    cur.execute("UPDATE contact_messages SET is_read=1 WHERE id=%s", (id,))
    mysql.connection.commit()
    cur.close()
    flash("Message marked as read.", "success")
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/delete_msg/<int:id>')
@admin_required
def delete_msg(id):
    cur = mysql.connection.cursor()
    cur.execute("DELETE FROM contact_messages WHERE id=%s", (id,))
    mysql.connection.commit()
    cur.close()
    flash("Message deleted.", "info")
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/unregister_student/<int:id>', methods=['POST'])
@admin_required
def unregister_student(id):
    cur = mysql.connection.cursor()
    try:
        cur.execute("DELETE FROM students WHERE id = %s", (id,))
        mysql.connection.commit()
        flash("Student unregistered successfully.", "success")
    except Exception as e:
        mysql.connection.rollback()
        logger.error(f"Unregister Error: {e}")
        flash("Error: Could not remove student.", "danger")
    finally:
        cur.close()
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/update_settings', methods=['POST'])
@admin_required
def update_settings():
    cur = mysql.connection.cursor()
    try:
        cur.execute("""
            UPDATE site_settings 
            SET office_location=%s, paybill_no=%s, paybill_acc_no=%s, bank_name=%s, 
                account_no=%s, account_name=%s, paypal_email=%s, mpesa_name=%s
            WHERE id=1
        """, (
            request.form.get('office_location'),
            request.form.get('paybill_no'),
            request.form.get('paybill_acc_no'),
            request.form.get('bank_name'),
            request.form.get('account_no'),
            request.form.get('account_name'),
            request.form.get('paypal_email'),
            request.form.get('mpesa_name')
        ))
        mysql.connection.commit()
        flash("Site settings updated.", "success")
    except Exception as e:
        mysql.connection.rollback()
        logger.error(f"Update error: {e}")
        flash("Error updating settings.", "danger")
    finally:
        cur.close()
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/export_students')
@admin_required
def export_students():
    campus_filter = request.args.get('campus', 'All').strip()
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    
    if campus_filter.lower() != "all" and campus_filter != "":
        cur.execute("SELECT full_name, email, phone, school_campus FROM students WHERE school_campus LIKE %s", (f"%{campus_filter}%",))
        title = f"Register: {campus_filter}"
    else:
        cur.execute("SELECT full_name, email, phone, school_campus FROM students")
        title = "Master Student Register"
    
    students = cur.fetchall()
    cur.close()

    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Name', 'Email', 'Phone', 'Campus'
    
    for s_info in students:
        row = table.add_row().cells
        row[0].text = str(s_info['full_name'])
        row[1].text = str(s_info['email'])
        row[2].text = str(s_info['phone'])
        row[3].text = str(s_info['school_campus'])

    target = BytesIO()
    doc.save(target)
    target.seek(0)
    return send_file(target, as_attachment=True, download_name=f"Students_{date.today()}.docx")

@app.route('/admin/add_event', methods=['POST'])
@admin_required
def add_event():
    title = request.form.get('title')
    desc = request.form.get('description')
    edate = request.form.get('event_date')
    loc = request.form.get('location')
    img_name = None
    if 'image' in request.files:
        f = request.files['image']
        if f and allowed_file(f.filename):
            img_name = f"media_{datetime.now().strftime('%Y%H%M%S')}_{secure_filename(f.filename)}"
            f.save(os.path.join(UPLOAD_FOLDER, img_name))
    cur = mysql.connection.cursor()
    cur.execute("INSERT INTO events (title, description, event_date, location, image) VALUES (%s, %s, %s, %s, %s)", 
                (title, desc, edate, loc, img_name))
    mysql.connection.commit()
    cur.close()
    flash("New event published.", "success")
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/edit_event/<int:id>', methods=['GET', 'POST'])
@admin_required
def edit_event(id):
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    if request.method == 'POST':
        title = request.form.get('title')
        desc = request.form.get('description')
        edate = request.form.get('event_date')
        loc = request.form.get('location')
        img_name = None
        if 'image' in request.files and request.files['image'].filename != '':
            f = request.files['image']
            if f and allowed_file(f.filename):
                cur.execute("SELECT image FROM events WHERE id=%s", (id,))
                old_data = cur.fetchone()
                if old_data and old_data['image']:
                    delete_physical_file(old_data['image'])
                img_name = f"media_{datetime.now().strftime('%Y%H%M%S')}_{secure_filename(f.filename)}"
                f.save(os.path.join(UPLOAD_FOLDER, img_name))
        
        if img_name:
            cur.execute("UPDATE events SET title=%s, description=%s, event_date=%s, location=%s, image=%s WHERE id=%s", 
                        (title, desc, edate, loc, img_name, id))
        else:
            cur.execute("UPDATE events SET title=%s, description=%s, event_date=%s, location=%s WHERE id=%s", 
                        (title, desc, edate, loc, id))
        mysql.connection.commit()
        cur.close()
        flash("Event updated.", "success")
        return redirect(url_for('admin_dashboard'))
        
    cur.execute("SELECT * FROM events WHERE id=%s", (id,))
    event = cur.fetchone()
    cur.close()
    return render_template('edit_event.html', event=event)

@app.route('/admin/delete_event/<int:id>')
@admin_required
def delete_event(id):
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cur.execute("SELECT image FROM events WHERE id=%s", (id,))
    event = cur.fetchone()
    if event and event['image']:
        delete_physical_file(event['image'])
    cur.execute("DELETE FROM events WHERE id=%s", (id,))
    mysql.connection.commit()
    cur.close()
    flash("Event removed.", "info")
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/upload_resource', methods=['POST'])
@admin_required
def upload_resource():
    name = request.form.get('display_name')
    f = request.files.get('resource_file')
    if f and allowed_file(f.filename):
        fname = secure_filename(f.filename)
        if os.path.exists(os.path.join(UPLOAD_FOLDER, fname)):
            fname = f"{datetime.now().strftime('%Y%H%M%S')}_{fname}"
        f.save(os.path.join(UPLOAD_FOLDER, fname))
        cur = mysql.connection.cursor()
        cur.execute("INSERT INTO downloads (display_name, filename) VALUES (%s, %s)", (name, fname))
        mysql.connection.commit()
        cur.close()
        flash("Resource uploaded.", "success")
    else:
        flash("Invalid file format.", "danger")
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/delete_resource/<int:id>')
@admin_required
def delete_resource(id):
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cur.execute("SELECT filename FROM downloads WHERE id=%s", (id,))
    res = cur.fetchone()
    if res and res['filename']:
        delete_physical_file(res['filename'])
    cur.execute("DELETE FROM downloads WHERE id=%s", (id,))
    mysql.connection.commit()
    cur.close()
    flash("Resource removed.", "info")
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/add_new_admin', methods=['POST'])
@admin_required
def add_new_admin():
    u = request.form.get('username')
    e = request.form.get('email')
    raw_p = request.form.get('password')
    confirm_p = request.form.get('confirm_password') # New field from HTML

    # 1. Basic Presence Check
    if not u or not e or not raw_p:
        flash("All fields (Username, Email, Password) are required.", "warning")
        return redirect(url_for('admin_dashboard'))

    # 2. Check if Passwords Match
    if raw_p != confirm_p:
        flash("Passwords do not match!", "danger")
        return redirect(url_for('admin_dashboard'))

    # 3. Validate Password Strength (Using the helper function)
    is_strong, message = is_strong_password(raw_p)
    if not is_strong:
        flash(message, "danger")
        return redirect(url_for('admin_dashboard'))

    # 4. Hash and Insert into Database
    p = generate_password_hash(raw_p)
    cur = mysql.connection.cursor()
    try:
        cur.execute("INSERT INTO admin (username, email, password) VALUES (%s, %s, %s)", (u, e, p))
        mysql.connection.commit()
        flash(f"New administrator {u} created successfully.", "success")
    except Exception as err:
        # It's better to log the error or check specifically for duplicate keys
        flash("Error: Username or Email already exists.", "danger")
    finally:
        cur.close()

    return redirect(url_for('admin_dashboard'))

@app.route('/admin/delete_admin/<int:id>')
@admin_required
def delete_admin(id):
    # 1. Fetch current admin ID from session
    current_admin_id = session.get('admin_id') 
    
    cur = mysql.connection.cursor()
    
    # 2. Perform the deletion (No "last admin" check anymore)
    # This allows you to wipe the table completely if needed
    cur.execute("DELETE FROM admin WHERE id=%s", (id,))
    mysql.connection.commit()
    cur.close()

    # 3. Logic for self-deletion
    if id == current_admin_id:
        session.clear()
        flash("Your account has been removed. The system is now empty or has other admins.", "info")
        # If this was the last admin, the next person to visit /admin/setup 
        # will see your green initialization screen.
        return redirect(url_for('admin_login'))

    flash("Administrator removed successfully.", "info")
    return redirect(url_for('admin_dashboard'))
@app.route('/admin/setup', methods=['GET', 'POST'])
def admin_setup():
    # 1. Database Check - Is the house empty?
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cur.execute("SELECT COUNT(*) as count FROM admin")
    admin_count = cur.fetchone()['count']
    
    # 2. Lockdown - Redirect if an admin already exists
    if admin_count > 0:
        cur.close()
        return redirect(url_for('admin_login'))

    if request.method == 'POST':
        # 3. Security Token Verification
        user_token = request.form.get('setup_token')
        if user_token != Config.SETUP_TOKEN:
            cur.close()
            flash("Unauthorized: Invalid Setup Token.", "danger")
            return render_template('admin_setup.html')

        # 4. Data Retrieval
        u = request.form.get('username')
        e = request.form.get('email')
        raw_p = request.form.get('password')
        confirm_p = request.form.get('confirm_password')

        # --- NEW SECURITY CHECKS ---
        # A. Check if Passwords Match
        if raw_p != confirm_p:
            cur.close()
            flash("Passwords do not match!", "danger")
            return render_template('admin_setup.html')

        # B. Validate Password Strength
        is_strong, message = is_strong_password(raw_p)
        if not is_strong:
            cur.close()
            flash(message, "danger")
            return render_template('admin_setup.html')
        # ---------------------------

        # 5. Master Admin Creation
        p = generate_password_hash(raw_p)
        try:
            cur.execute("INSERT INTO admin (username, email, password) VALUES (%s, %s, %s)", (u, e, p))
            mysql.connection.commit()
            flash("System initialized. Master account created!", "success")
            return redirect(url_for('admin_login'))
        except Exception as err:
            flash(f"Initialization Error: {str(err)}", "danger")
        finally:
            cur.close()

    return render_template('admin_setup.html')
# --- NEW GALLERY ROUTES ---

@app.route('/gallery')
def gallery():
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    # Fetching from 'gallery' table
    cur.execute("SELECT * FROM gallery ORDER BY created_at DESC")
    images = cur.fetchall() # Changed variable from 'photos' to 'images'
    cur.close()
    # Passing 'images' to match the loop in gallery.html
    return render_template('gallery.html', images=images)

@app.route('/admin/add_gallery', methods=['POST'])
@admin_required
def add_gallery():
    # We fetch the text. If it's empty, Flask just gives us ""
    title = request.form.get('title', '') 
    category = request.form.get('category', '')
    f = request.files.get('photo')

    # The ONLY thing we must have is the file
    if not f:
        flash("Please select a photo.", "danger")
        return redirect(url_for('admin_dashboard'))

    if f and allowed_file(f.filename):
        fname = f"gallery_{datetime.now().strftime('%Y%H%M%S')}_{secure_filename(f.filename)}"
        f.save(os.path.join(UPLOAD_FOLDER, fname))
        
        cur = mysql.connection.cursor()
        # It's okay if title or category are blank strings here
        cur.execute("INSERT INTO gallery (title, image_path, category) VALUES (%s, %s, %s)", 
                    (title, fname, category))
        mysql.connection.commit()
        cur.close()
        flash("Uploaded successfully!", "success")
    else:
        flash("Invalid file format.", "danger")
        
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/delete_gallery/<int:id>')
@admin_required
def delete_gallery(id):
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cur.execute("SELECT image_path FROM gallery WHERE id=%s", (id,))
    photo = cur.fetchone()
    if photo:
        if photo['image_path']:
            # Ensure this function is defined to remove the file from static/uploads
            delete_physical_file(photo['image_path'])
        cur.execute("DELETE FROM gallery WHERE id=%s", (id,))
        mysql.connection.commit()
        flash("Photo deleted.", "info")
    cur.close()
    return redirect(url_for('admin_dashboard'))
from flask import send_from_directory

@app.route('/download_photo/<filename>')
def download_photo(filename):
    return send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True)
# --- 1. PUBLIC VIEW ---
# --- 1. PUBLIC VIEW ---
# --- 1. PUBLIC VIEW: LIST ALUMNI ---
@app.route('/alumni')
def alumni_view():
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cur.execute("SELECT * FROM alumni ORDER BY grad_year DESC")
    alumni_data = cur.fetchall()
    cur.close()
    return render_template('alumni.html', alumni=alumni_data)
@app.route('/admin/dashboard')
@admin_required 
def admin_dashboard():
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    
    # Fetch all necessary data
    cur.execute("SELECT * FROM events ORDER BY event_date DESC")
    ev = cur.fetchall()
    
    cur.execute("SELECT * FROM students ORDER BY registered_at DESC")
    st = cur.fetchall()
    
    cur.execute("SELECT * FROM downloads ORDER BY uploaded_at DESC")
    dl = cur.fetchall()
    
    cur.execute("SELECT id, username, email FROM admin")
    adm = cur.fetchall()
    
    cur.execute("SELECT * FROM site_settings WHERE id=1")
    finance_data = cur.fetchone()
    
    cur.execute("SELECT * FROM contact_messages ORDER BY is_read ASC, created_at DESC")
    msgs = cur.fetchall()
    
    cur.execute("SELECT * FROM gallery ORDER BY created_at DESC")
    gallery_items = cur.fetchall()

    cur.execute("SELECT * FROM alumni ORDER BY created_at DESC")
    alumni_list = cur.fetchall()

    cur.execute("SELECT COUNT(*) as total FROM alumni")
    alumni_count = cur.fetchone()['total']
    
    cur.close()
    
    return render_template('admin_dashboard.html', 
                            events=ev, students=st, downloads=dl, 
                            admins=adm, today=date.today(),
                            finance=finance_data, messages=msgs,
                            gallery_items=gallery_items,
                            alumni=alumni_list,
                            alumni_count=alumni_count)

# --- 2. PUBLIC REGISTRATION ---
@app.route('/register_alumni', methods=['GET', 'POST'])
def register_alumni():
    if request.method == 'POST':
        full_name = request.form.get('name')
        email = request.form.get('email')
        grad_year = request.form.get('grad_year')
        course = request.form.get('course')
        location = request.form.get('location')
        profession = request.form.get('profession')
        
        cur = mysql.connection.cursor()
        try:
            cur.execute("""
                INSERT INTO alumni (full_name, email, grad_year, course_studied, location, profession) 
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (full_name, email, grad_year, course, location, profession))
            mysql.connection.commit()
            flash('Successfully joined the Alumni Network!', 'success')
        except Exception as e:
            mysql.connection.rollback()
            flash(f'Error: {str(e)}', 'danger')
        finally:
            cur.close()
        
        return redirect(url_for('index')) 
        
    return render_template('alumni_register.html')

# --- 3. ADMIN: MANAGE ALUMNI ---
@app.route('/admin/manage_alumni')
def admin_manage_alumni():
    if not session.get('loggedin'):
        return redirect(url_for('login'))
        
    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cur.execute("SELECT * FROM alumni ORDER BY created_at DESC")
    all_alumni = cur.fetchall()
    cur.close()
    
    return render_template('admin_dashboard.html', alumni=all_alumni)

# --- 4. ADMIN: DELETE ALUMNI ---
@app.route('/admin/delete_alumni/<int:id>')
def delete_alumni(id):
    if not session.get('loggedin'):
        return redirect(url_for('login'))
        
    cur = mysql.connection.cursor()
    cur.execute("DELETE FROM alumni WHERE id = %s", [id])
    mysql.connection.commit()
    cur.close()
    
    flash('Alumni record removed.', 'info')
    return redirect(url_for('admin_dashboard'))

# --- 5. ADMIN: EXPORT ALUMNI TO WORD (.DOCX) ---
@app.route('/admin/export_alumni')
def export_alumni():
    if not session.get('loggedin'):
        return redirect(url_for('login'))

    cur = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
    cur.execute("SELECT full_name, email, grad_year, course_studied, profession FROM alumni")
    alumni_list = cur.fetchall()
    cur.close()

    # Create Word Document
    doc = Document()
    
    # Line 910: Fixed Title Alignment
    title = doc.add_heading('Lutheran Campus Ministry', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Line 913: Fixed Subtitle Alignment
    subtitle = doc.add_paragraph('Official Alumni Network Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Date Generated: {date.today().strftime('%B %d, %Y')}")
    doc.add_paragraph("_" * 50)

    # Add Table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Header Row
    hdr_cells = table.rows[0].cells
    headers = ['Full Name', 'Email', 'Year', 'Course', 'Profession']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        # Make headers bold
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.runs[0]
        run.bold = True

    # Add Data Rows
    for person in alumni_list:
        row_cells = table.add_row().cells
        row_cells[0].text = str(person['full_name'] or '')
        row_cells[1].text = str(person['email'] or '')
        row_cells[2].text = str(person['grad_year'] or '')
        row_cells[3].text = str(person['course_studied'] or '')
        row_cells[4].text = str(person['profession'] or '')

    # Line 940: Fixed Byte Stream Handling
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name='LCM_Alumni_Report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
@app.route('/logout')
def logout():
    session.clear()
    flash("Logged out.", "info")
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(host='127.0.0.1', port=5000, debug=True)